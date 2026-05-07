import os
import json

from resume_tailorator.models.agents.output import FinalReport
from resume_tailorator.models.workflow import ResumeTailorResult
from resume_tailorator.utils.pdf_converter import markdown_to_pdf


def generate_resume(
    result: ResumeTailorResult, output_dir: str, base_filename: str
) -> str:
    """
    Convert tailored CV to Markdown, PDF, and DOCX formats.

    Args:
        result: ResumeTailorResult object containing tailored resume and company name
        output_dir: Directory to save output files (job-specific subdirectory).
        base_filename: Base name for all output files (without extension).

    Returns:
        The path to the generated Markdown file.
    """
    os.makedirs(output_dir, exist_ok=True)
    # Validate base_filename doesn't contain path separators or parent refs
    if os.sep in base_filename or (os.altsep and os.altsep in base_filename):
        raise ValueError(f"base_filename contains path separator: {base_filename}")
    if ".." in base_filename:
        raise ValueError(f"base_filename contains parent-directory reference: {base_filename}")
    md_output_path = os.path.join(output_dir, f"{base_filename}.md")
    pdf_output_path = os.path.join(output_dir, f"{base_filename}.pdf")
    docx_output_path = os.path.join(output_dir, f"{base_filename}.docx")

    # Parse the CV JSON back to CV object
    cv_data = json.loads(result.tailored_resume)

    # Build markdown content
    md_content = [f"# {cv_data.get('full_name', 'N/A')}\n"]

    if cv_data.get("contact_info"):
        md_content.append(f"{cv_data.get('contact_info')}\n")
    md_content.append("\n")

    md_content.append(f"## Professional Summary\n{cv_data.get('summary', '')}\n\n")

    md_content.append("## Skills\n")
    for skill in cv_data.get("skills", []):
        md_content.append(f"- {skill}\n")

    if cv_data.get("projects"):
        md_content.append("\n## Projects\n")
        for project in cv_data.get("projects", []):
            md_content.append(f"{project}\n\n")

    md_content.append("\n## Work Experience\n")
    for exp in cv_data.get("experience", []):
        md_content.append(
            f"### {exp.get('role', '')} at {exp.get('company', '')} ({exp.get('dates', '')})\n\n"
        )
        for hl in exp.get("highlights", []):
            md_content.append(f"- {hl}\n")
        md_content.append("\n")

    md_content.append("## Education\n")
    for edu in cv_data.get("education", []):
        md_content.append(f"- {edu}\n")

    if cv_data.get("certifications"):
        md_content.append("\n## Certifications\n")
        for cert in cv_data.get("certifications", []):
            md_content.append(f"- {cert}\n")

    if cv_data.get("publications"):
        md_content.append("\n## Publications\n")
        for pub in cv_data.get("publications", []):
            md_content.append(f"- {pub}\n")

    markdown_text = "".join(md_content)

    # Save Markdown
    with open(md_output_path, "w", encoding="utf-8") as f:
        f.write(markdown_text)

    # Save PDF
    markdown_to_pdf(markdown_text, pdf_output_path)

    # Save DOCX
    try:
        from docx import Document

        doc = Document()
        for line in markdown_text.split("\n"):
            if line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:].strip(), style="List Bullet")
            elif line.strip():
                doc.add_paragraph(line.strip())
        doc.save(docx_output_path)
    except Exception as e:
        print(f"⚠️ Warning: Failed to save DOCX: {e}")

    print(
        f"✅ Tailored CV saved to:\n   - Markdown: {md_output_path}\n   - PDF: {pdf_output_path}\n   - DOCX: {docx_output_path}"
    )

    return md_output_path


def generate_report_markdown(report: FinalReport) -> str:
    """Render a FinalReport as a Markdown string.

    Args:
        report: The completed FinalReport.

    Returns:
        A Markdown-formatted string ready to write to a file.
    """
    lines: list[str] = []

    lines.append(f"# Self-Review Report — {report.company_name} · {report.job_title}\n")
    lines.append(f"**Generated:** {report.generated_at}  \n")
    lines.append(f"**Audit Passed:** {'✅ Yes' if report.passed else '❌ No'}\n")

    lines.append("---\n")

    # Match score and recommendation
    lines.append("## 🎯 Match Score & Recommendation\n")
    lines.append(f"**Score:** {report.match_score}/100  \n")
    lines.append(f"**Verdict:** {report.overall_recommendation}\n")
    lines.append(f"{report.recommendation_rationale}\n")

    # What changed
    lines.append("---\n")
    lines.append("## ✏️ What Changed\n")
    diff = report.what_changed
    if not diff.sections_modified:
        lines.append("_No significant changes detected._\n")
    else:
        rendered_any = False
        if diff.summary_changed:
            lines.append("- **Summary** was rewritten\n")
            rendered_any = True
        if diff.skills_reordered:
            reordered = ", ".join(diff.skills_reordered)
            lines.append(f"- **Skills reordered to top:** {reordered}\n")
            rendered_any = True
        if diff.skills_deprioritized:
            deprioritized = ", ".join(diff.skills_deprioritized)
            lines.append(f"- **Skills deprioritized:** {deprioritized}\n")
            rendered_any = True
        for exp_change in diff.experience_changes:
            lines.append(
                f"- **{exp_change.role} at {exp_change.company}:** "
                f"{len(exp_change.bullets_rephrased)} bullet(s) rephrased, "
                f"{exp_change.bullets_unchanged} unchanged\n"
            )
            for bullet in exp_change.bullets_rephrased:
                lines.append(f"  - {bullet}\n")
            rendered_any = True
        if not rendered_any:
            lines.append("_Changes noted but details unavailable._\n")

    # Keyword coverage
    lines.append("---\n")
    lines.append("## 🔑 Keyword Coverage\n")
    gap = report.gaps
    total = len(gap.covered_keywords) + len(gap.missing_keywords)
    if total > 0:
        pct = (len(gap.covered_keywords) / total) * 100
        coverage_str = f"{len(gap.covered_keywords)}/{total}"
    else:
        pct = 0.0
        coverage_str = "N/A"
    lines.append(f"**Keywords covered: {coverage_str} ({pct:.1f}%)**\n")
    if gap.covered_keywords:
        covered_str = ", ".join(f"`{k}`" for k in gap.covered_keywords)
        lines.append(f"\n✅ **Covered:** {covered_str}\n")
    if gap.missing_keywords:
        missing_str = ", ".join(f"`{k}`" for k in gap.missing_keywords)
        lines.append(f"\n❌ **Missing:** {missing_str}\n")

    # Skill gaps
    lines.append("---\n")
    lines.append("## 🚧 Skill Gaps\n")
    if not gap.missing_hard_skills and not gap.missing_soft_skills:
        lines.append("_No skill gaps detected — your CV covers all required skills!_\n")
    else:
        if gap.missing_hard_skills:
            hard_str = ", ".join(gap.missing_hard_skills)
            lines.append(f"**Hard skills not in your CV:** {hard_str}\n")
        if gap.missing_soft_skills:
            soft_str = ", ".join(gap.missing_soft_skills)
            lines.append(f"**Soft skills not in your CV:** {soft_str}\n")

    # Suggestions
    lines.append("---\n")
    lines.append("## 💡 Suggestions to Strengthen Your Application\n")
    if not report.suggestions_to_strengthen:
        lines.append("_No additional suggestions — your application looks strong!_\n")
    else:
        for suggestion in report.suggestions_to_strengthen:
            lines.append(f"- {suggestion}\n")

    # Audit summary
    lines.append("---\n")
    lines.append("## 🔍 Audit Summary\n")
    lines.append(f"{report.audit_summary}\n")

    return "".join(lines)
