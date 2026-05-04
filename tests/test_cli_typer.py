"""Tests for CLI with Typer - tailor and re-tailor commands."""

from unittest.mock import AsyncMock, MagicMock, patch
import os

import pytest
from typer.testing import CliRunner

from resume_tailorator.main import app
from resume_tailorator.models.agents.output import CV, WorkExperience, ScrapedJobPosting
from resume_tailorator.models.workflow import ResumeTailorResult

runner = CliRunner()


def _make_cv(full_name: str = "Jane Doe") -> CV:
    return CV(
        full_name=full_name,
        contact_info="jane@example.com",
        summary="Platform engineer.",
        skills=["Python", "SQL"],
        experience=[
            WorkExperience(
                company="Acme",
                role="Engineer",
                dates="2022-2026",
                highlights=["Built services"],
            )
        ],
        education=["BSc CS"],
    )


def _make_result(cv: CV | None = None, passed: bool = True) -> ResumeTailorResult:
    cv = cv or _make_cv()
    from resume_tailorator.models.agents.output import CVDiff, FinalReport, GapAnalysis
    final_report = FinalReport(
        company_name="Acme Corp",
        job_title="Software Engineer",
        generated_at="2026-01-01T00:00:00Z",
        overall_recommendation="Strong Match",
        match_score=85,
        what_changed=CVDiff(sections_modified=["summary"]),
        gaps=GapAnalysis(
            covered_keywords=["Python"],
            missing_keywords=["Rust"],
        ),
        suggestions_to_strengthen=[ "Add Rust experience"],
        audit_summary="Looks good.",
        recommendation_rationale="Strong alignment with role requirements.",
        passed=passed,
    )
    return ResumeTailorResult(
        company_name="Acme Corp",
        job_title="Software Engineer",
        tailored_resume=cv.model_dump_json(),
        audit_report={
            "passed": passed,
            "hallucination_score": 0,
            "ai_cliche_score": 1,
            "issues": [],
            "feedback_summary": "Looks good.",
        },
        passed=passed,
        final_report=final_report,
    )


def _make_scraped_job(markdown: str = "# Job Posting\nPython engineer") -> ScrapedJobPosting:
    return ScrapedJobPosting(
        markdown=markdown,
        url="https://example.com/job/123",
        source_text="Raw job posting text",
        extraction_strategy="test",
    )


def test_tailor_command_success(tmp_path, monkeypatch) -> None:
    """tailor command with valid inputs should succeed and show output."""
    resume_file = tmp_path / "resume.md"
    resume_file.write_text("# Jane Doe\nPython developer.")

    output_dir = tmp_path / "output"
    output_dir.mkdir()

    monkeypatch.chdir(tmp_path)

    cv = _make_cv()
    workflow_result = _make_result(cv=cv, passed=True)

    mock_workflow = MagicMock()
    mock_workflow.run = AsyncMock(return_value=workflow_result)
    mock_generate_resume = MagicMock(return_value=str(output_dir / "tailored_resume_acme_corp.md"))

    scraped_job = _make_scraped_job()

    with (
        patch(
            "resume_tailorator.main.job_scraper_agent.run",
            AsyncMock(return_value=MagicMock(output=scraped_job)),
        ),
        patch("resume_tailorator.main.ResumeTailorWorkflow", return_value=mock_workflow),
        patch("resume_tailorator.main.generate_resume", mock_generate_resume),
        patch("resume_tailorator.main.SQLiteResumeMemoryRepository") as mock_repo_cls,
        patch("resume_tailorator.main.PydanticAIResumeParser") as mock_parser_cls,
        patch("resume_tailorator.main.ResumeMemoryService") as mock_svc_cls,
    ):
        mock_repo = MagicMock()
        mock_repo_cls.return_value = mock_repo
        mock_parser = MagicMock()
        mock_parser_cls.return_value = mock_parser
        mock_svc = MagicMock()
        mock_svc.resolve_original_resume.return_value = MagicMock(
            source=MagicMock(id="src-123"),
            cv=cv,
        )
        mock_svc.save_tailored_resume.return_value = MagicMock(id="job-456")
        mock_svc_cls.return_value = mock_svc

        result = runner.invoke(
            app,
            [
                "tailor",
                "https://example.com/job/123",
                str(resume_file),
                "--output-dir",
                str(output_dir),
            ],
        )

    assert result.exit_code == 0, result.output
    assert "✅ Job completed" in result.output
    mock_workflow.run.assert_called_once()
    mock_generate_resume.assert_called_once()


def test_tailor_command_invalid_url_format(tmp_path, monkeypatch) -> None:
    """tailor command with invalid URL format should return 1."""
    resume_file = tmp_path / "resume.md"
    resume_file.write_text("# Jane Doe\nPython developer.")

    output_dir = tmp_path / "output"
    output_dir.mkdir()

    monkeypatch.chdir(tmp_path)

    result = runner.invoke(
        app,
        [
            "tailor",
            "not-a-valid-url",
            str(resume_file),
            "--output-dir",
            str(output_dir),
        ],
    )

    assert result.exit_code == 1
    assert "http://" in result.output or "https://" in result.output


def test_tailor_command_resume_not_found(tmp_path, monkeypatch) -> None:
    """tailor command with non-existent resume should return 1."""
    output_dir = tmp_path / "output"
    output_dir.mkdir()

    monkeypatch.chdir(tmp_path)

    result = runner.invoke(
        app,
        [
            "tailor",
            "https://example.com/job/123",
            "/no/such/file.md",
            "--output-dir",
            str(output_dir),
        ],
    )

    assert result.exit_code == 1
    assert "not found" in result.output.lower() or "❌" in result.output


def test_tailor_command_scraping_failure(tmp_path, monkeypatch) -> None:
    """tailor command when scraping fails should return 1."""
    resume_file = tmp_path / "resume.md"
    resume_file.write_text("# Jane Doe\nPython developer.")

    output_dir = tmp_path / "output"
    output_dir.mkdir()

    monkeypatch.chdir(tmp_path)

    mock_scraper = AsyncMock(side_effect=Exception("Network error"))

    with patch("resume_tailorator.main.job_scraper_agent.run", mock_scraper):
        result = runner.invoke(
            app,
            [
                "tailor",
                "https://example.com/job/123",
                str(resume_file),
                "--output-dir",
                str(output_dir),
            ],
        )

    assert result.exit_code == 1
    assert "scrape" in result.output.lower() or "failed" in result.output.lower()


def test_tailor_command_failed_audit_exits_zero(tmp_path, monkeypatch) -> None:
    """tailor command with failed audit should exit 0 (report still generated)."""
    resume_file = tmp_path / "resume.md"
    resume_file.write_text("# Jane Doe\nPython developer.")

    output_dir = tmp_path / "output"
    output_dir.mkdir()

    monkeypatch.chdir(tmp_path)

    cv = _make_cv()
    workflow_result = _make_result(cv=cv, passed=False)

    mock_workflow = MagicMock()
    mock_workflow.run = AsyncMock(return_value=workflow_result)
    mock_generate_resume = MagicMock(return_value=str(output_dir / "tailored_resume_acme_corp.md"))

    scraped_job = _make_scraped_job()

    with (
        patch(
            "resume_tailorator.main.job_scraper_agent.run",
            AsyncMock(return_value=MagicMock(output=scraped_job)),
        ),
        patch("resume_tailorator.main.ResumeTailorWorkflow", return_value=mock_workflow),
        patch("resume_tailorator.main.generate_resume", mock_generate_resume),
        patch("resume_tailorator.main.SQLiteResumeMemoryRepository") as mock_repo_cls,
        patch("resume_tailorator.main.PydanticAIResumeParser") as mock_parser_cls,
        patch("resume_tailorator.main.ResumeMemoryService") as mock_svc_cls,
    ):
        mock_repo = MagicMock()
        mock_repo_cls.return_value = mock_repo
        mock_parser = MagicMock()
        mock_parser_cls.return_value = mock_parser
        mock_svc = MagicMock()
        mock_svc.resolve_original_resume.return_value = MagicMock(
            source=MagicMock(id="src-123"),
            cv=cv,
        )
        mock_svc.save_tailored_resume.return_value = MagicMock(id="job-456")
        mock_svc_cls.return_value = mock_svc

        result = runner.invoke(
            app,
            [
                "tailor",
                "https://example.com/job/123",
                str(resume_file),
                "--output-dir",
                str(output_dir),
            ],
        )

    assert result.exit_code == 0, result.output
    mock_workflow.run.assert_called_once()


def test_tailor_command_empty_job_content(tmp_path, monkeypatch) -> None:
    """tailor command when scraped job is empty should return 1."""
    resume_file = tmp_path / "resume.md"
    resume_file.write_text("# Jane Doe\nPython developer.")

    output_dir = tmp_path / "output"
    output_dir.mkdir()

    monkeypatch.chdir(tmp_path)

    scraped_job = _make_scraped_job(markdown="")

    with patch(
        "resume_tailorator.main.job_scraper_agent.run",
        AsyncMock(return_value=MagicMock(output=scraped_job)),
    ):
        result = runner.invoke(
            app,
            [
                "tailor",
                "https://example.com/job/123",
                str(resume_file),
                "--output-dir",
                str(output_dir),
            ],
        )

    assert result.exit_code == 1
    assert "empty" in result.output.lower()


def test_tailor_command_docx_conversion(tmp_path, monkeypatch) -> None:
    """tailor command should convert DOCX resume to markdown."""
    # Use docx library to create a minimal docx file.
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    resume_file = tmp_path / "resume.docx"
    doc = Document()
    doc.add_heading("Jane Doe", 0)
    doc.add_paragraph("Python developer")
    doc.save(str(resume_file))

    output_dir = tmp_path / "output"
    output_dir.mkdir()

    monkeypatch.chdir(tmp_path)

    cv = _make_cv()
    workflow_result = _make_result(cv=cv, passed=True)

    mock_workflow = MagicMock()
    mock_workflow.run = AsyncMock(return_value=workflow_result)
    mock_generate_resume = MagicMock(return_value=str(output_dir / "tailored_resume_acme_corp.md"))

    scraped_job = _make_scraped_job()

    with (
        patch(
            "resume_tailorator.main.job_scraper_agent.run",
            AsyncMock(return_value=MagicMock(output=scraped_job)),
        ),
        patch("resume_tailorator.main.ResumeTailorWorkflow", return_value=mock_workflow),
        patch("resume_tailorator.main.generate_resume", mock_generate_resume),
        patch("resume_tailorator.main.SQLiteResumeMemoryRepository") as mock_repo_cls,
        patch("resume_tailorator.main.PydanticAIResumeParser") as mock_parser_cls,
        patch("resume_tailorator.main.ResumeMemoryService") as mock_svc_cls,
    ):
        mock_repo = MagicMock()
        mock_repo_cls.return_value = mock_repo
        mock_parser = MagicMock()
        mock_parser_cls.return_value = mock_parser
        mock_svc = MagicMock()
        mock_svc.resolve_original_resume.return_value = MagicMock(
            source=MagicMock(id="src-123"),
            cv=cv,
        )
        mock_svc.save_tailored_resume.return_value = MagicMock(id="job-456")
        mock_svc_cls.return_value = mock_svc

        result = runner.invoke(
            app,
            [
                "tailor",
                "https://example.com/job/123",
                str(resume_file),
                "--output-dir",
                str(output_dir),
            ],
        )

    assert result.exit_code == 0, result.output
    mock_workflow.run.assert_called_once()
    converted_file = output_dir / "resume_converted.md"
    assert converted_file.exists()


# --- re_tailor tests ---


def test_re_tailor_success(tmp_path, monkeypatch) -> None:
    """re_tailor with valid job ID and recommendations should succeed."""
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    monkeypatch.chdir(tmp_path)

    cv = _make_cv()
    workflow_result = _make_result(cv=cv, passed=True)

    mock_workflow = MagicMock()
    mock_workflow.run = AsyncMock(return_value=workflow_result)
    mock_generate_resume = MagicMock(return_value=str(output_dir / "tailored_resume_acme_corp.md"))

    with (
        patch("resume_tailorator.main.ResumeTailorWorkflow", return_value=mock_workflow),
        patch("resume_tailorator.main.generate_resume", mock_generate_resume),
        patch("resume_tailorator.main.SQLiteResumeMemoryRepository") as mock_repo_cls,
        patch("resume_tailorator.main.PydanticAIResumeParser") as mock_parser_cls,
        patch("resume_tailorator.main.ResumeMemoryService") as mock_svc_cls,
    ):
        mock_repo = MagicMock()
        mock_repo.get_tailored_resume_by_id.return_value = MagicMock(
            source_id="src-123",
            company_name="Acme Corp",
            job_title="Software Engineer",
            job_fingerprint="fp123",
            job_posting_markdown="# Job Posting\nPython engineer",
        )
        mock_repo.get_source_by_id.return_value = MagicMock(path="/resume.md")
        mock_repo_cls.return_value = mock_repo

        mock_parser = MagicMock()
        mock_parser_cls.return_value = mock_parser

        mock_svc = MagicMock()
        mock_svc.resolve_original_resume.return_value = MagicMock(
            source=MagicMock(id="src-123"),
            cv=cv,
        )
        mock_svc_cls.return_value = mock_svc

        result = runner.invoke(
            app,
            [
                "re-tailor",
                "job-456",
                "Add more detail about leadership skills",
                "--output-dir",
                str(output_dir),
            ],
        )

    assert result.exit_code == 0, result.output
    assert "✅ Re-tailoring completed" in result.output
    mock_workflow.run.assert_called_once()
    mock_repo.save_tailored_resume.assert_called_once()


def test_re_tailor_job_not_found(tmp_path, monkeypatch) -> None:
    """re_tailor with non-existent job ID should return 1."""
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    monkeypatch.chdir(tmp_path)

    with patch("resume_tailorator.main.SQLiteResumeMemoryRepository") as mock_repo_cls:
        mock_repo = MagicMock()
        mock_repo.get_tailored_resume_by_id.return_value = None
        mock_repo_cls.return_value = mock_repo

        result = runner.invoke(
            app,
            [
                "re-tailor",
                "nonexistent-id",
                "Add leadership skills",
                "--output-dir",
                str(output_dir),
            ],
        )

    assert result.exit_code == 1
    assert "Job not found" in result.output


def test_re_tailor_no_job_markdown(tmp_path, monkeypatch) -> None:
    """re_tailor should fail when stored job has no posting markdown."""
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    monkeypatch.chdir(tmp_path)

    with patch("resume_tailorator.main.SQLiteResumeMemoryRepository") as mock_repo_cls:
        mock_repo = MagicMock()
        mock_repo.get_tailored_resume_by_id.return_value = MagicMock(
            source_id="src-123",
            company_name="Acme Corp",
            job_title="Software Engineer",
            job_fingerprint="fp123",
            job_posting_markdown="",
        )
        mock_repo.get_source_by_id.return_value = MagicMock(path="/resume.md")
        mock_repo_cls.return_value = mock_repo

        with (
            patch("resume_tailorator.main.PydanticAIResumeParser") as mock_parser_cls,
            patch("resume_tailorator.main.ResumeMemoryService") as mock_svc_cls,
        ):
            mock_svc = MagicMock()
            mock_svc.resolve_original_resume.return_value = MagicMock(
                source=MagicMock(id="src-123"),
                cv=_make_cv(),
            )
            mock_svc_cls.return_value = mock_svc

            result = runner.invoke(
                app,
                [
                    "re-tailor",
                    "job-456",
                    "Add leadership skills",
                    "--output-dir",
                    str(output_dir),
                ],
            )

    assert result.exit_code == 1
    assert "No job posting content stored" in result.output


def test_re_tailor_with_resume_path(tmp_path, monkeypatch) -> None:
    """re_tailor should use explicit resume path when provided."""
    resume_file = tmp_path / "resume.md"
    resume_file.write_text("# Jane Doe\nPython developer.")
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    monkeypatch.chdir(tmp_path)

    cv = _make_cv()
    workflow_result = _make_result(cv=cv, passed=True)

    mock_workflow = MagicMock()
    mock_workflow.run = AsyncMock(return_value=workflow_result)
    mock_generate_resume = MagicMock(return_value=str(output_dir / "tailored_resume_acme_corp.md"))

    with (
        patch("resume_tailorator.main.ResumeTailorWorkflow", return_value=mock_workflow),
        patch("resume_tailorator.main.generate_resume", mock_generate_resume),
        patch("resume_tailorator.main.SQLiteResumeMemoryRepository") as mock_repo_cls,
        patch("resume_tailorator.main.PydanticAIResumeParser") as mock_parser_cls,
        patch("resume_tailorator.main.ResumeMemoryService") as mock_svc_cls,
    ):
        mock_repo = MagicMock()
        mock_repo.get_tailored_resume_by_id.return_value = MagicMock(
            source_id="src-123",
            company_name="Acme Corp",
            job_title="Software Engineer",
            job_fingerprint="fp123",
            job_posting_markdown="# Job Posting\nPython engineer",
        )
        mock_repo_cls.return_value = mock_repo

        mock_svc = MagicMock()
        mock_svc.resolve_original_resume.return_value = MagicMock(
            source=MagicMock(id="src-123"),
            cv=cv,
        )
        mock_svc_cls.return_value = mock_svc

        result = runner.invoke(
            app,
            [
                "re-tailor",
                "job-456",
                "Add leadership skills",
                "--resume-path",
                str(resume_file),
                "--output-dir",
                str(output_dir),
            ],
        )

    assert result.exit_code == 0, result.output
    # When --resume-path is passed, we should not call get_source_by_id.
    mock_repo.get_source_by_id.assert_not_called()
    mock_svc.resolve_original_resume.assert_called_with(path=str(resume_file))


def test_re_tailor_missing_resume_file(tmp_path, monkeypatch) -> None:
    """re_tailor should fail when explicit resume path does not exist."""
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    monkeypatch.chdir(tmp_path)

    with patch("resume_tailorator.main.SQLiteResumeMemoryRepository") as mock_repo_cls:
        mock_repo = MagicMock()
        mock_repo.get_tailored_resume_by_id.return_value = MagicMock(
            source_id="src-123",
            company_name="Acme Corp",
            job_title="Software Engineer",
            job_fingerprint="fp123",
            job_posting_markdown="# Job Posting",
        )
        mock_repo_cls.return_value = mock_repo

        result = runner.invoke(
            app,
            [
                "re-tailor",
                "job-456",
                "Add skills",
                "--resume-path",
                "/no/such/file.md",
                "--output-dir",
                str(output_dir),
            ],
        )

    assert result.exit_code == 1
    assert "not found" in result.output.lower() or "❌" in result.output
